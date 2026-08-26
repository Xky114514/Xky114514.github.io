from __future__ import annotations

from pathlib import Path
from typing import Iterable

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips


ROOT = Path(r"D:\guanmai\vibe_coding\Xky114514.github.io")
WORK = ROOT / ".manual_work"
SHOTS = WORK / "screenshots"
OUTPUT = Path(r"D:\guanmai\vibe_coding\采购录单应用客户操作手册_V1.0.docx")
FLOW_IMAGE = WORK / "purchase_workflow.png"

INK = "243047"
NAVY = "17365D"
BLUE = "2E74B5"
BLUE_DARK = "1F4D78"
MUTED = "667085"
LIGHT_BLUE = "E8F1FB"
LIGHT_GRAY = "F5F7FA"
LINE = "D8E0EA"
WHITE = "FFFFFF"
AMBER = "FFF4D6"
AMBER_TEXT = "7A5200"
GREEN = "EAF7E5"
GREEN_TEXT = "2E6B24"
RED = "FCE8E6"
RED_TEXT = "9B1C1C"

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}
FIGURE_WIDTH = Inches(6.25)


def ensure_child(parent, tag: str):
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.append(child)
    return child


def set_run_font(run, *, size: float | None = None, bold: bool | None = None,
                 color: str | None = None, italic: bool | None = None,
                 latin: str = "Calibri", east_asia: str = "Microsoft YaHei") -> None:
    run.font.name = latin
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), latin)
    r_fonts.set(qn("w:hAnsi"), latin)
    r_fonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if italic is not None:
        run.italic = italic


def style_paragraph(paragraph, *, before: float = 0, after: float = 6,
                    line_spacing: float = 1.25, keep_with_next: bool = False,
                    keep_together: bool = False) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line_spacing
    fmt.keep_with_next = keep_with_next
    fmt.keep_together = keep_together
    fmt.widow_control = True


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = ensure_child(tc_pr, "w:tcMar")
    for side, margin in CELL_MARGINS_DXA.items():
        node = ensure_child(tc_mar, f"w:{side}")
        node.set(qn("w:type"), "dxa")
        node.set(qn("w:w"), str(margin))


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def apply_table_geometry(table, widths: list[int], *, indent: int = TABLE_INDENT_DXA) -> None:
    if sum(widths) != CONTENT_WIDTH_DXA:
        raise ValueError(f"table widths must sum to {CONTENT_WIDTH_DXA}, got {sum(widths)}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = ensure_child(tbl_pr, "w:tblW")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_ind = ensure_child(tbl_pr, "w:tblInd")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(indent))
    layout = ensure_child(tbl_pr, "w:tblLayout")
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        prevent_row_split(row)
        for column_index, cell in enumerate(row.cells):
            width = widths[column_index]
            cell.width = Twips(width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = ensure_child(tc_pr, "w:tcW")
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))
            set_cell_margins(cell)


def set_table_borders(table, color: str = LINE, *, size: int = 5) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for child in list(borders):
        borders.remove(child)
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge = OxmlElement(f"w:{name}")
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), str(size))
        edge.set(qn("w:space"), "0")
        edge.set(qn("w:color"), color)
        borders.append(edge)


def create_numbering(doc: Document) -> tuple[int, int]:
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    next_abs = max(abstract_ids, default=0) + 1
    next_num = max(num_ids, default=0) + 1

    def add_definition(abstract_id: int, num_id: int, num_fmt: str, lvl_text: str) -> None:
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        level = OxmlElement("w:lvl")
        level.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        fmt = OxmlElement("w:numFmt")
        fmt.set(qn("w:val"), num_fmt)
        text = OxmlElement("w:lvlText")
        text.set(qn("w:val"), lvl_text)
        suffix = OxmlElement("w:suff")
        suffix.set(qn("w:val"), "tab")
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "540")
        tabs.append(tab)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "540")
        ind.set(qn("w:hanging"), "270")
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "300")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.extend([tabs, ind, spacing])
        level.extend([start, fmt, text, suffix, p_pr])
        abstract.append(level)
        numbering.append(abstract)
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abs_ref = OxmlElement("w:abstractNumId")
        abs_ref.set(qn("w:val"), str(abstract_id))
        num.append(abs_ref)
        numbering.append(num)

    add_definition(next_abs, next_num, "bullet", "•")
    add_definition(next_abs + 1, next_num + 1, "decimal", "%1.")
    return next_num, next_num + 1


def set_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])


def add_body(doc, text: str, *, after: float = 6, color: str = INK,
             bold_lead: str | None = None, italic: bool = False) -> None:
    p = doc.add_paragraph()
    style_paragraph(p, after=after)
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run_font(lead, size=11, bold=True, color=NAVY)
        rest = p.add_run(text[len(bold_lead):])
        set_run_font(rest, size=11, color=color, italic=italic)
    else:
        run = p.add_run(text)
        set_run_font(run, size=11, color=color, italic=italic)


def add_bullet(doc, text: str, bullet_num_id: int, *, bold_lead: str | None = None) -> None:
    p = doc.add_paragraph()
    set_numbering(p, bullet_num_id)
    style_paragraph(p, after=4)
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run_font(lead, size=11, bold=True, color=NAVY)
        rest = p.add_run(text[len(bold_lead):])
        set_run_font(rest, size=11, color=INK)
    else:
        run = p.add_run(text)
        set_run_font(run, size=11, color=INK)


def add_step(doc, step_no: int, title: str, detail: str) -> None:
    p = doc.add_paragraph()
    style_paragraph(p, before=6, after=2, line_spacing=1.15, keep_with_next=True)
    badge = p.add_run(f"步骤 {step_no}  ")
    set_run_font(badge, size=10.5, bold=True, color=WHITE)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), BLUE)
    badge._element.get_or_add_rPr().append(shd)
    title_run = p.add_run(title)
    set_run_font(title_run, size=11.5, bold=True, color=NAVY)
    d = doc.add_paragraph()
    style_paragraph(d, after=5, line_spacing=1.25)
    d.paragraph_format.left_indent = Inches(0.18)
    r = d.add_run(detail)
    set_run_font(r, size=10.5, color=INK)


def add_heading(doc, text: str, level: int = 1) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    for run in p.runs:
        set_run_font(run)


def add_callout(doc, label: str, text: str, *, fill: str = LIGHT_BLUE,
                accent: str = BLUE, text_color: str = INK) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_repeat_table_header(table.rows[0])
    apply_table_geometry(table, [CONTENT_WIDTH_DXA])
    set_table_borders(table, accent, size=7)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    style_paragraph(p, after=0, line_spacing=1.25)
    lead = p.add_run(f"{label}  ")
    set_run_font(lead, size=10.5, bold=True, color=accent)
    body = p.add_run(text)
    set_run_font(body, size=10.5, color=text_color)
    spacer = doc.add_paragraph()
    style_paragraph(spacer, after=2, line_spacing=1)


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[int],
              *, center_cols: Iterable[int] = ()) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_repeat_table_header(table.rows[0])
    center_set = set(center_cols)
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_paragraph(p, after=0, line_spacing=1.1)
        run = p.add_run(text)
        set_run_font(run, size=9.3, bold=True, color=NAVY)
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for column_index, value in enumerate(values):
            if row_index % 2:
                set_cell_shading(cells[column_index], LIGHT_GRAY)
            p = cells[column_index].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if column_index in center_set else WD_ALIGN_PARAGRAPH.LEFT
            style_paragraph(p, after=0, line_spacing=1.12)
            run = p.add_run(value)
            set_run_font(run, size=9.2, color=INK, bold=(column_index == 0))
    apply_table_geometry(table, widths)
    set_table_borders(table)
    spacer = doc.add_paragraph()
    style_paragraph(spacer, after=3, line_spacing=1)


def set_image_alt(inline_shape, title: str, description: str) -> None:
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("title", title)
    doc_pr.set("descr", description)


def add_figure(doc, filename: str, caption: str, *, width=FIGURE_WIDTH) -> None:
    path = SHOTS / filename
    if not path.exists():
        raise FileNotFoundError(path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(p, before=4, after=3, line_spacing=1, keep_with_next=True, keep_together=True)
    shape = p.add_run().add_picture(str(path), width=width)
    set_image_alt(shape, caption, f"采购录单应用界面截图：{caption}")
    cap = doc.add_paragraph(style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(cap, after=8, line_spacing=1.1, keep_together=True)
    run = cap.add_run(caption)
    set_run_font(run, size=9, color=MUTED, italic=True)


def add_flow_figure(doc, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(p, before=4, after=3, line_spacing=1, keep_with_next=True, keep_together=True)
    shape = p.add_run().add_picture(str(FLOW_IMAGE), width=FIGURE_WIDTH)
    set_image_alt(shape, caption, "采购单与采购入库单两条业务链路流程图")
    cap = doc.add_paragraph(style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(cap, after=8, line_spacing=1.1)
    run = cap.add_run(caption)
    set_run_font(run, size=9, color=MUTED, italic=True)


def add_page_break(doc) -> None:
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_page_field(paragraph, field_name: str) -> None:
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_name
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, sep, text, end])
    set_run_font(run, size=9, color=MUTED)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, BLUE_DARK, 10, 5),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def configure_page(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style_paragraph(hp, after=0, line_spacing=1)
    left = hp.add_run("采购录单应用")
    set_run_font(left, size=9, bold=True, color=NAVY)
    right = hp.add_run("  |  客户操作手册")
    set_run_font(right, size=9, color=MUTED)
    p_pr = hp._p.get_or_add_pPr()
    borders = ensure_child(p_pr, "w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "5")
    bottom.set(qn("w:space"), "5")
    bottom.set(qn("w:color"), LINE)
    borders.append(bottom)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    style_paragraph(fp, after=0, line_spacing=1)
    lead = fp.add_run("V1.0  |  第 ")
    set_run_font(lead, size=9, color=MUTED)
    add_page_field(fp, "PAGE")
    tail = fp.add_run(" 页")
    set_run_font(tail, size=9, color=MUTED)


def make_workflow_image() -> None:
    canvas = Image.new("RGB", (1800, 760), "white")
    draw = ImageDraw.Draw(canvas)
    font_path = r"C:\Windows\Fonts\msyh.ttc"
    bold_path = r"C:\Windows\Fonts\msyhbd.ttc"
    font = ImageFont.truetype(font_path, 34)
    small = ImageFont.truetype(font_path, 28)
    bold = ImageFont.truetype(bold_path, 42)
    label = ImageFont.truetype(bold_path, 34)

    draw.rounded_rectangle((20, 20, 1780, 740), radius=34, fill="#F8FAFD", outline="#D8E0EA", width=3)
    draw.text((70, 54), "采购录单应用端到端流程", font=bold, fill="#17365D")

    lanes = [
        ("采购单链路", 175, "#E8F1FB", "#2E74B5", [
            "选择供应商/采购群", "发送文字或附件", "AI 生成采购单", "人工核对明细", "确认提交观麦", "观麦审核"
        ]),
        ("采购入库单链路", 435, "#EAF7E5", "#3A7D2C", [
            "选择供应商/入库群", "发送到货信息", "AI 生成入库草稿", "关联采购单并核对", "确认/补单", "观麦审核"
        ]),
    ]
    for lane_name, y, fill, accent, steps in lanes:
        draw.rounded_rectangle((60, y, 285, y + 155), radius=24, fill=accent)
        bbox = draw.multiline_textbbox((0, 0), lane_name.replace("链路", "\n链路"), font=label, spacing=8, align="center")
        tw = bbox[2] - bbox[0]
        th = bbox[3] - bbox[1]
        draw.multiline_text((172 - tw / 2, y + 77 - th / 2), lane_name.replace("链路", "\n链路"), font=label, fill="white", spacing=8, align="center")
        x = 330
        box_w = 210
        gap = 38
        for i, step in enumerate(steps):
            draw.rounded_rectangle((x, y + 18, x + box_w, y + 137), radius=18, fill=fill, outline=accent, width=3)
            bbox = draw.multiline_textbbox((0, 0), step.replace("/", "/\n"), font=small, spacing=5, align="center")
            tw = bbox[2] - bbox[0]
            th = bbox[3] - bbox[1]
            draw.multiline_text((x + box_w / 2 - tw / 2, y + 78 - th / 2), step.replace("/", "/\n"), font=small, fill="#243047", spacing=5, align="center")
            if i < len(steps) - 1:
                ax = x + box_w + 8
                ay = y + 78
                draw.line((ax, ay, ax + gap - 14, ay), fill=accent, width=5)
                draw.polygon([(ax + gap - 14, ay), (ax + gap - 28, ay - 10), (ax + gap - 28, ay + 10)], fill=accent)
            x += box_w + gap
    canvas.save(FLOW_IMAGE, quality=95)


def build() -> None:
    make_workflow_image()
    doc = Document()
    doc.core_properties.title = "采购录单应用客户操作手册"
    doc.core_properties.subject = "采购单与采购入库单端到端操作流程"
    doc.core_properties.author = "采购录单项目组"
    doc.core_properties.keywords = "采购录单,采购单,采购入库单,客户操作手册"
    configure_styles(doc)
    configure_page(doc)
    bullet_num_id, _decimal_num_id = create_numbering(doc)

    # Cover: editorial_cover pattern.
    spacer = doc.add_paragraph()
    style_paragraph(spacer, after=0)
    spacer.paragraph_format.space_before = Pt(72)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(kicker, after=16, line_spacing=1)
    kr = kicker.add_run("CUSTOMER OPERATION GUIDE  ·  V1.0")
    set_run_font(kr, size=10.5, bold=True, color=BLUE)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(title, after=8, line_spacing=1.05, keep_with_next=True)
    tr = title.add_run("采购录单应用\n客户操作手册")
    set_run_font(tr, size=30, bold=True, color=NAVY)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(subtitle, after=24, line_spacing=1.15)
    sr = subtitle.add_run("覆盖采购单录入、采购入库单录入、审核确认、补单及基础配置")
    set_run_font(sr, size=13.5, color=BLUE_DARK)
    add_figure(doc, "01-purchase-home.png", "采购录单工作台（示例数据）", width=Inches(5.9))
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(meta, before=8, after=2, line_spacing=1.2)
    mr = meta.add_run("适用版本：V1.0    |    更新日期：2026-08-05    |    适用对象：采购操作员、审核人员、客户管理员")
    set_run_font(mr, size=9.5, color=MUTED)

    add_page_break(doc)

    add_heading(doc, "使用说明", 1)
    add_callout(doc, "阅读提示", "本文从已登录后的采购首页开始说明。账号开通、登录地址和权限分配由客户管理员按实际部署环境提供；当前原型未单独展示登录页面。界面中的供应商、群聊、单号和数量均为示例数据。")
    add_body(doc, "本手册按“先理解流程、再完成录单、最后处理异常”的顺序编排。建议首次使用者先阅读第 1～2 章，再按采购单或采购入库单的实际业务选择对应章节操作。")

    add_heading(doc, "目录", 2)
    add_table(doc, ["章节", "主要内容"], [
        ["1 业务范围与角色", "两条单据链路、角色分工、操作前准备"],
        ["2 首页与通用操作", "快捷入口、任务状态、筛选与来源定位"],
        ["3 采购单操作流程", "录入、AI 识别、审核、明细核对、提交"],
        ["4 采购入库单操作流程", "录入、关联采购单、数量金额核对、确认与补单"],
        ["5 基础配置与统计", "供应商、群聊、提示词、AI 记忆、机器人、统计"],
        ["6 异常处理与检查清单", "常见异常、边界规则、提交前/后检查"],
    ], [1900, 7460])

    add_heading(doc, "1 业务范围与角色", 1)
    add_heading(doc, "1.1 两条独立业务链路", 2)
    add_body(doc, "采购录单应用同时承载采购计划和实际到货两类业务，但两类单据必须分开录入、分开审核，避免把“计划采购数量”和“实际入库数量”混为一谈。")
    add_flow_figure(doc, "图 1  采购单与采购入库单端到端流程")
    add_callout(doc, "关键边界", "采购单录入只生成采购单；采购入库单录入只生成采购入库单草稿。系统不会因为录入采购单而自动生成入库单，也不会在录入阶段自动合并两类任务。", fill=AMBER, accent=AMBER_TEXT, text_color=AMBER_TEXT)

    add_heading(doc, "1.2 角色分工", 2)
    add_table(doc, ["角色", "主要职责", "典型功能"], [
        ["采购操作员", "录入采购/到货信息，核对 AI 识别结果，保存或提交单据", "采购单录入、采购入库单录入、审核列表、详情"],
        ["审核/复核人员", "检查供应商、商品、数量、价格、关联关系和异常备注", "审核筛选、详情核对、重识别、确认提交"],
        ["客户管理员", "维护供应商、群聊、提示词、AI 记忆和机器人配置", "供应商、群聊管理、知识配置、设置"],
        ["观麦操作员", "在观麦系统中完成最终业务审核；审核后单据不可继续补充", "观麦采购单/采购入库单审核"],
    ], [1550, 4100, 3710])

    add_heading(doc, "1.3 操作前准备", 2)
    for text in [
        "确认当前账号已开通采购录单应用，并具有对应供应商或群聊的处理权限。",
        "确认供应商主数据已经同步；从群聊录入时，确认群聊已绑定正确供应商和单据类型。",
        "确认机器人在线，且群聊处于允许采购发言/采集的时间段。",
        "准备清晰的采购文字、手写单、图片、微信截图、Excel 或 PDF；一条消息尽量只描述同一供应商、同一类单据。",
        "涉及采购入库时，提前核实采购单是否已在观麦同步，以及该采购单是否仍处于可关联状态。",
    ]:
        add_bullet(doc, text, bullet_num_id)

    add_page_break(doc)
    add_heading(doc, "2 首页与通用操作", 1)
    add_heading(doc, "2.1 认识采购首页", 2)
    add_body(doc, "采购首页是日常工作台。顶部快捷入口用于进入各功能；“今日任务”汇总采购单和采购入库单的处理压力；下方同步看板用于判断采购单是否仍可关联或补单。")
    add_figure(doc, "01-purchase-home.png", "图 2  采购首页：快捷入口与今日任务")
    add_step(doc, 1, "选择业务入口", "计划采购进入“采购单录入/审核”；供应商实际到货进入“采购入库单录入/审核”。")
    add_step(doc, 2, "关注待处理数量", "优先处理“待处理”和“失败”任务；“已提交”表示单据已传到观麦等待后续审核。")
    add_step(doc, 3, "检查同步提示", "采购单已审核或关闭时，不能继续关联或补充入库内容；应改选可用采购单或独立提交入库单。")

    add_heading(doc, "2.2 常用状态说明", 2)
    add_table(doc, ["状态", "业务含义", "建议操作"], [
        ["待处理", "AI 已生成任务，但尚未提交观麦", "进入详情核对并保存或确认提交"],
        ["暂存", "当前修改已保存，但仍未提交观麦", "后续继续核对；确认无误再提交"],
        ["已提交", "已提交观麦，等待观麦侧审核", "通常只读；到观麦跟进最终审核"],
        ["失败", "识别或回传未成功", "检查原始内容，使用重识别或重新录入"],
        ["已关闭/已审核", "观麦侧业务已结束", "不可继续关联或补单"],
    ], [1500, 3900, 3960])

    add_heading(doc, "2.3 列表页通用操作", 2)
    for text in [
        "查询：设置状态、日期范围、群聊、操作员或供应商后点击“查询”。",
        "重置：清空筛选条件并恢复默认列表。",
        "刷新：重新读取当前任务状态。",
        "查看：进入单据详情，核对原始消息和 AI 识别明细。",
        "重识别：识别结果明显偏离原文时使用；重识别后仍需人工复核。",
        "删除：仅处理尚未提交的错误任务；已提交单据不应在录单应用重复删除。",
    ]:
        add_bullet(doc, text, bullet_num_id, bold_lead=text.split("：")[0] + "：")

    add_page_break(doc)
    add_heading(doc, "3 采购单操作流程", 1)
    add_heading(doc, "3.1 新建采购单", 2)
    add_body(doc, "采购单用于记录计划采购内容，例如次日采购品项、数量和送达要求。录入时先确定供应商或采购单群，再发送文字或附件。")
    add_figure(doc, "02-purchase-order-entry.png", "图 3  采购单录入初始界面")
    add_step(doc, 1, "选择来源", "在左侧选择“供应商”或“群聊”。供应商模式适合单独录入；群聊模式适合按已绑定采购群处理。")
    add_step(doc, 2, "定位供应商/群", "可通过名称或 ID 搜索。必须先选中来源，右侧输入区才具备明确归属。")
    add_step(doc, 3, "输入或上传采购内容", "可发送文字，也可拖拽、粘贴或上传图片、Excel、PDF。原文应尽量包含商品、数量/单位、采购日期和备注。")
    add_figure(doc, "02c-purchase-order-filled.png", "图 4  已选择供应商并填写采购内容")
    add_step(doc, 4, "点击发送", "系统调用 AI 识别商品和采购数量，并生成待处理采购单任务。收到生成成功提示后，到“采购单审核”继续处理。")
    add_callout(doc, "录入建议", "同一条消息不要混入多个供应商或采购入库信息；价格缺失可留待人工确认，但商品名称和数量/单位必须能够从原文判断。", fill=AMBER, accent=AMBER_TEXT, text_color=AMBER_TEXT)

    add_heading(doc, "3.2 在审核列表找到任务", 2)
    add_figure(doc, "03-purchase-order-review.png", "图 5  采购单审核列表")
    add_step(doc, 1, "设置筛选条件", "按状态、采购时间、群聊、操作员和供应商筛选。新生成任务通常处于“待处理”。")
    add_step(doc, 2, "核对原文摘要", "确认列表中的群聊、供应商、原文和商品数与本次录入一致。")
    add_step(doc, 3, "分配操作员并查看", "如需转交，可调整操作员；点击对应行“查看”进入详情。")

    add_page_break(doc)
    add_heading(doc, "3.3 核对采购单详情", 2)
    add_figure(doc, "04-purchase-order-detail.png", "图 6  采购单详情与 AI 识别明细")
    add_body(doc, "详情页左侧用于追溯原始文件和群聊消息，右侧用于修正 AI 识别结果。操作时应逐行对照，不要只看识别后的表格。")
    add_step(doc, 1, "核对来源", "查看原始文件、原始消息、供应商、来源群聊和采购日期，必要时使用“定位来源”。")
    add_step(doc, 2, "核对供应商和备注", "供应商应与原始采购要求一致；采购备注用于保留价格待确认、临采、交付要求等信息。")
    add_step(doc, 3, "逐行核对商品", "重点检查商品名称、数量/单位、商品备注、SPU 名称和商品编码。识别错误可直接修改，也可新增/删除商品行。")
    add_step(doc, 4, "保存或确认提交", "“保存”只保留当前修改；“确认提交”会把采购单传至观麦，状态变为已提交。")
    add_figure(doc, "04b-purchase-order-confirm.png", "图 7  采购单提交成功提示")
    add_callout(doc, "提交前务必确认", "采购单确认提交后将进入观麦审核流程。请先确认供应商、数量/单位和商品映射；不要依赖提交后的再次修改。", fill=RED, accent=RED_TEXT, text_color=RED_TEXT)

    add_page_break(doc)
    add_heading(doc, "4 采购入库单操作流程", 1)
    add_heading(doc, "4.1 新建采购入库单", 2)
    add_body(doc, "采购入库单用于记录供应商实际到货。应使用实收数量、入库单价和异常备注，不应直接照搬采购计划数量。")
    add_figure(doc, "05-purchase-entry.png", "图 8  采购入库单录入界面")
    add_step(doc, 1, "选择供应商或入库群", "群聊模式只选择已标记为“采购入库单录入”的群；不要选择采购单群。")
    add_step(doc, 2, "发送到货信息", "内容建议包含到货商品、实际数量/单位、到货/入库时间、单价及破损、短缺、复称等异常。")
    add_step(doc, 3, "生成待处理草稿", "AI 识别到货数和价格后生成采购入库单草稿；随后进入“采购入库单审核”。")

    add_heading(doc, "4.2 在审核列表选择任务", 2)
    add_figure(doc, "06-purchase-review.png", "图 9  采购入库单审核列表")
    add_body(doc, "审核列表展示状态、送货情况、入库时间、来源群聊、供应商、原文、商品数和操作员。送货情况用于帮助评审和判断分批场景，不能替代详情核对。")
    add_step(doc, 1, "筛选待处理任务", "按状态、入库时间、群聊、操作员和供应商查询。")
    add_step(doc, 2, "识别分批/未关联场景", "列表可能显示“一单一入库”“分批送货”“暂未关联采购单”等提示，进入详情后再决定关联或独立提交。")
    add_step(doc, 3, "点击查看", "在详情页完成关联采购单、数量金额核对、保存和确认提交。")

    add_page_break(doc)
    add_heading(doc, "4.3 核对来源并关联采购单", 2)
    add_figure(doc, "07-purchase-inbound-detail.png", "图 10  入库详情：来源、关联采购单与基本信息")
    add_step(doc, 1, "追溯原始到货信息", "先查看左侧原始文件和群聊消息，确认供应商、到货时间和分批描述。")
    add_step(doc, 2, "选择采购单时间范围", "默认查看近 1 天；找不到采购单时可切换近 3 天、近 7 天或全部时间。")
    add_step(doc, 3, "关联一张采购单或暂不关联", "每张采购入库单最多关联一张仍可处理的采购单；采购单尚未同步时可以不关联并继续提交。")
    add_step(doc, 4, "查看历史入库批次", "“查看关联采购入库单”用于判断是否已有历史批次、暂存或待确认单据；查看不会自动合并或提交。")
    add_callout(doc, "供应商联动规则", "供应商在“AI 识别采购入库单明细”区域编辑，关联采购单区域只读同步。若更换供应商导致原采购单不匹配，应先确认解除或改选关联。", fill=AMBER, accent=AMBER_TEXT, text_color=AMBER_TEXT)

    add_heading(doc, "4.4 核对商品、实收数和金额", 2)
    add_figure(doc, "07b-purchase-inbound-lines.png", "图 11  入库明细：采购数量、实收数与差异提示")
    add_table(doc, ["字段", "核对要求"], [
        ["识别文本", "原始消息中与本行对应的内容，用于人工追溯"],
        ["商品名称/编码", "应映射到正确 SPU；名称相近时仍需核对编码"],
        ["采购数量", "来自关联采购单；未关联采购单时可为空"],
        ["实收数", "本次实际收到的数量，是入库的核心数量"],
        ["差异", "采购数量与实收数的差额；严格超过采购数量 10% 时高亮，但不阻断提交"],
        ["单价/金额", "金额默认按实收数×单价计算；如人工改过金额，后续数量或单价变化可能不再自动覆盖"],
        ["备注", "记录复称、短缺、破损、分批、价格异常等可追溯信息"],
    ], [1900, 7460])
    add_step(doc, 1, "逐行对照原文", "确认商品、单位和实收数；不要用采购数量代替实收数。")
    add_step(doc, 2, "处理差异与价格", "对高亮差异和缺失/异常价格进行人工确认，并在备注中说明原因。")
    add_step(doc, 3, "必要时增删商品", "AI 漏识别时新增商品；重复或错误行可删除。已提交单据不允许继续增删。")
    add_step(doc, 4, "保存当前结果", "尚未确认或需要等待复称时点击“保存”；保存不会提交观麦，也不会联动其他批次。")

    add_page_break(doc)
    add_heading(doc, "4.5 确认提交与补单分支", 2)
    add_body(doc, "点击“确认提交”后，系统会检查是否存在可补单的观麦待处理入库单。不同检查结果对应不同操作路径。")
    add_heading(doc, "情形 A：没有可补单记录", 3)
    add_step(doc, 1, "复核当前单", "确认供应商、入库时间、商品数、实收数、价格和备注无误。")
    add_step(doc, 2, "确认提交", "在确认提示中再次点击“确认提交”，当前采购入库单将独立提交观麦。")
    add_callout(doc, "允许未关联提交", "采购单未及时同步时，采购入库单仍可独立提交；请在备注中保留“采购单待同步”等说明，便于后续追溯。", fill=GREEN, accent=GREEN_TEXT, text_color=GREEN_TEXT)

    add_heading(doc, "情形 B：存在可补单记录", 3)
    add_figure(doc, "08-supplement-candidates.png", "图 12  系统发现可补单的入库单")
    add_body(doc, "候选记录通常同时满足：供应商一致、同一入库日期、候选时间不晚于当前单、已提交观麦且尚未审核。系统给出候选不代表必须补单，操作员仍需核对业务事实。")
    add_step(doc, 1, "查看候选摘要", "核对候选单号、时间和商品数，避免把不同送货业务误补到一起。")
    add_step(doc, 2, "查看详情", "打开候选商品详情，核对实收数量、单价、金额和备注。")
    add_figure(doc, "09-supplement-detail.png", "图 13  可补单候选的商品详情")
    add_step(doc, 3, "选择处理方式", "属于同一张观麦待处理入库单时点击“补充到此单”；应独立成单时点击弹窗底部“确认提交”。")
    add_step(doc, 4, "完成二次确认", "补充到候选单会再次确认；成功后到观麦核对被补充单的商品明细。")
    add_callout(doc, "不可补单的情形", "候选单或关联采购单已在观麦审核/关闭后，不能继续补充。此时应返回详情解除或更换关联，并按独立入库单处理。", fill=RED, accent=RED_TEXT, text_color=RED_TEXT)

    add_page_break(doc)
    add_heading(doc, "5 基础配置与统计", 1)
    add_heading(doc, "5.1 供应商管理", 2)
    add_figure(doc, "10-suppliers.png", "图 14  供应商管理")
    add_body(doc, "供应商主数据决定录入归属、采购单可关联范围和 AI 记忆积累。客户管理员应定期检查供应商 ID、名称、联系人、地址、电话、群绑定和 SKU 同步数量。")
    for text in [
        "通过供应商名称、ID 或地址搜索，并按绑定状态筛选。",
        "未绑定群聊的供应商不能稳定从群聊识别归属，应先完成群绑定。",
        "SKU 同步数为 0 或明显异常时，先检查主数据同步，再处理识别偏差。",
        "供应商重名时必须依赖唯一 ID，不要仅按显示名称判断。",
    ]:
        add_bullet(doc, text, bullet_num_id)

    add_heading(doc, "5.2 采购群聊管理", 2)
    add_figure(doc, "11-groups.png", "图 15  采购群聊管理")
    add_body(doc, "群聊类型决定消息进入采购单链路还是采购入库单链路。绑定错误会直接造成任务生成到错误入口。")
    add_step(doc, 1, "检查群聊类型", "“采购单录入”用于计划采购；“采购入库单录入”用于实际到货。")
    add_step(doc, 2, "检查供应商与操作员", "确认绑定供应商数量和责任操作员符合业务分工。")
    add_step(doc, 3, "检查机器人与时段", "机器人应处于正常发言/在线状态，采购时段应覆盖实际消息时间。")

    add_heading(doc, "5.3 提示词管理", 2)
    add_figure(doc, "12-prompts.png", "图 16  采购提示词管理")
    add_body(doc, "采购单和采购入库单使用不同提示词页签；每类又可分供应商提示词和系统提示词。修改模板会影响后续识别，应先在小范围样例验证。")
    for text in [
        "采购单提示词聚焦采购计划、商品和采购数量。",
        "采购入库单提示词聚焦实际到货、实收数、单价和入库备注。",
        "启用前检查绑定供应商和模板类型，避免一个模板覆盖不相关供应商。",
        "保留更新时间和变更原因，出现识别回退时便于追溯。",
    ]:
        add_bullet(doc, text, bullet_num_id)

    add_page_break(doc)
    add_heading(doc, "5.4 AI 记忆", 2)
    add_figure(doc, "13-memory.png", "图 17  采购 AI 记忆")
    add_body(doc, "AI 记忆按采购单/采购入库单、供应商/群聊隔离，记录商品修正、下单习惯、常购商品和累计订单。")
    add_step(doc, 1, "定位对象", "选择正确单据类型，再按供应商或群聊查看。")
    add_step(doc, 2, "查看详情", "核对历史修正是否仍符合当前商品和业务口径。")
    add_step(doc, 3, "谨慎清理", "错误记忆会持续影响识别；清理前确认不是偶发输入问题，并保留变更记录。")

    add_heading(doc, "5.5 机器人与计算设置", 2)
    add_figure(doc, "14-settings.png", "图 18  采购设置：机器人管理")
    add_body(doc, "设置页包含机器人管理、回复设置和计算优先级。日常重点检查机器人在线状态；计算优先级变更会影响数量/金额口径，应由客户管理员统一维护。")
    add_callout(doc, "权限建议", "提示词、AI 记忆和计算优先级属于高影响配置，不建议普通操作员随意修改。变更前应准备测试样例，变更后抽检新生成任务。", fill=AMBER, accent=AMBER_TEXT, text_color=AMBER_TEXT)

    add_heading(doc, "5.6 统计与导出", 2)
    add_figure(doc, "15-statistics.png", "图 19  采购录单统计")
    add_body(doc, "统计页按时间范围汇总操作员下单数、提交数和商品总数，用于工作量复盘和运营对账。可选择今日、昨日、本周、本月、上月或自定义日期，并导出 Excel。")
    add_bullet(doc, "统计反映录单处理量，不等同于观麦最终审核通过量。", bullet_num_id)
    add_bullet(doc, "导出前确认日期范围和时区，避免跨日数据遗漏。", bullet_num_id)

    add_page_break(doc)
    add_heading(doc, "6 异常处理与检查清单", 1)
    add_heading(doc, "6.1 常见异常处理", 2)
    add_table(doc, ["现象", "可能原因", "处理建议"], [
        ["找不到供应商", "供应商未同步、名称不一致或账号无权限", "用供应商 ID 搜索；检查供应商主数据和权限"],
        ["群聊不出现在录入页", "群聊类型错误、未绑定或被禁言", "到群聊管理检查类型、供应商、操作员和状态"],
        ["AI 识别商品/数量错误", "原图模糊、单位缺失、提示词或记忆偏差", "对照原文人工修改；必要时重识别并检查配置"],
        ["采购入库单找不到采购单", "供应商不同、时间范围过窄、采购单未同步或已关闭", "扩大时间范围；核对供应商；允许暂不关联提交"],
        ["差异被红色高亮", "实收数与采购数量差异严格超过 10%", "核实实收数并在备注说明；高亮不阻断提交"],
        ["无法继续补单", "目标单已在观麦审核或关闭", "独立提交当前入库单，或改选仍可处理的目标单"],
        ["价格或金额异常", "价格缺失、人工金额被锁定或单位不一致", "核对单价与单位；必要时直接修正金额并备注"],
        ["任务显示失败", "识别失败、回传异常或原始文件不可读", "查看原始内容，重识别；仍失败则重新录入并联系管理员"],
    ], [2100, 3300, 3960])

    add_heading(doc, "6.2 提交前检查清单", 2)
    for text in [
        "单据类型正确：采购计划进入采购单，实际到货进入采购入库单。",
        "供应商、来源群聊和操作员与实际业务一致。",
        "已逐行对照原文，商品名称、编码、数量和单位无误。",
        "采购入库单的实收数、单价、金额和入库时间已经确认。",
        "数量差异、短缺、破损、复称、分批和价格异常均已写入备注。",
        "关联采购单仍处于可处理状态；若未关联，已确认可独立提交。",
        "发现补单候选时已查看详情，确认补单或独立提交的业务归属。",
    ]:
        add_bullet(doc, text, bullet_num_id)

    add_heading(doc, "6.3 提交后检查清单", 2)
    for text in [
        "确认录单应用出现提交成功提示，任务状态变为“已提交”。",
        "到观麦系统核对单据是否到达、供应商和商品明细是否完整。",
        "补单场景重点核对被补充单的新增商品、数量、价格和金额。",
        "通知观麦操作员及时审核；审核后不再继续补充或修改。",
        "如发现错误，立即记录单号、来源消息和错误字段，并按客户内部纠错流程处理。",
    ]:
        add_bullet(doc, text, bullet_num_id)

    add_heading(doc, "6.4 核心规则速查", 2)
    add_callout(doc, "规则 1", "每张采购入库单最多关联一张采购单，也可以暂不关联。")
    add_callout(doc, "规则 2", "一张采购单可对应一张或多张采购入库单，适用于分批送货。")
    add_callout(doc, "规则 3", "“保存”只保存当前修改；“确认提交”才会把单据送到观麦。")
    add_callout(doc, "规则 4", "已提交且尚未在观麦审核的入库单可能成为补单候选；已审核/关闭单据不可补。")
    add_callout(doc, "规则 5", "采购数量与实收数差异严格超过 10% 时高亮，但系统仍允许提交，责任人必须人工判断并备注。")

    add_heading(doc, "文档结束", 1)
    add_body(doc, "如实际部署版本与本文截图存在差异，请以客户管理员发布的版本说明和当前系统页面为准。建议在每次提示词、AI 记忆或计算规则调整后更新本手册版本。", color=MUTED, italic=True)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
