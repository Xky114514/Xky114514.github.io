from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips


OUTPUT = Path(__file__).resolve().parent / "assets" / "微信订单机器人部署方案及收费标准_商务排版优化版.docx"

INK = "243047"
NAVY = "17365D"
BLUE = "2E74B5"
BLUE_DARK = "1F4D78"
MUTED = "667085"
LIGHT_BLUE = "E8F1FB"
LIGHT_GRAY = "F5F7FA"
LINE = "D8E0EA"
WHITE = "FFFFFF"
AMBER = "FFF4D6"
AMBER_TEXT = "7A5200"

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS_DXA = {"top": 100, "bottom": 100, "start": 120, "end": 120}


def set_run_font(run, *, size: float | None = None, bold: bool | None = None,
                 color: str | None = None, italic: bool | None = None) -> None:
    run.font.name = "Arial"
    run._element.get_or_add_rPr().get_or_add_rFonts()
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if italic is not None:
        run.italic = italic


def style_paragraph(paragraph, *, before: float = 0, after: float = 6,
                    line_spacing: float = 1.25, keep_with_next: bool = False) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line_spacing
    fmt.keep_with_next = keep_with_next


def shade_paragraph(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_borders(paragraph, *, color: str = LINE, left: bool = False) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    if left:
        side = OxmlElement("w:left")
        side.set(qn("w:val"), "single")
        side.set(qn("w:sz"), "18")
        side.set(qn("w:space"), "10")
        side.set(qn("w:color"), color)
        borders.append(side)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def _ensure_child(parent, tag: str):
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.append(child)
    return child


def apply_table_geometry(table, widths: list[int]) -> None:
    if sum(widths) != CONTENT_WIDTH_DXA:
        raise ValueError(f"table widths must sum to {CONTENT_WIDTH_DXA}, got {sum(widths)}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = _ensure_child(tbl_pr, "w:tblW")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))

    tbl_ind = _ensure_child(tbl_pr, "w:tblInd")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))

    layout = _ensure_child(tbl_pr, "w:tblLayout")
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for column_index, width in enumerate(widths):
        table.columns[column_index].width = Twips(width)
    for row in table.rows:
        prevent_row_split(row)
        for column_index, cell in enumerate(row.cells):
            width = widths[column_index]
            cell.width = Twips(width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = _ensure_child(tc_pr, "w:tcW")
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))
            tc_mar = _ensure_child(tc_pr, "w:tcMar")
            for side, margin in CELL_MARGINS_DXA.items():
                node = _ensure_child(tc_mar, f"w:{side}")
                node.set(qn("w:type"), "dxa")
                node.set(qn("w:w"), str(margin))


def set_table_borders(table, color: str = LINE) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge = OxmlElement(f"w:{name}")
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), "5")
        edge.set(qn("w:space"), "0")
        edge.set(qn("w:color"), color)
        borders.append(edge)


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[int],
              *, alignments: list[int] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.rows[0].height = None
    set_repeat_table_header(table.rows[0])
    for index, text in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, LIGHT_BLUE)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_paragraph(paragraph, after=0, line_spacing=1.1)
        run = paragraph.add_run(text)
        set_run_font(run, size=9.5, bold=True, color=NAVY)

    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for column_index, value in enumerate(values):
            cell = cells[column_index]
            if row_index % 2:
                set_cell_shading(cell, LIGHT_GRAY)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = (
                alignments[column_index]
                if alignments
                else WD_ALIGN_PARAGRAPH.LEFT
            )
            style_paragraph(paragraph, after=0, line_spacing=1.08)
            run = paragraph.add_run(value)
            set_run_font(run, size=9.3, color=INK, bold=(column_index == 0))

    apply_table_geometry(table, widths)
    set_table_borders(table)
    spacer = doc.add_paragraph()
    style_paragraph(spacer, after=2, line_spacing=1)


def add_heading(doc, text: str, level: int = 1) -> None:
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    run = paragraph.add_run(text)
    if not paragraph.text.startswith(text):
        paragraph.text = text
    for item in paragraph.runs:
        set_run_font(item)


def add_body(doc, text: str, *, bold_lead: str | None = None,
             color: str = INK, after: float = 6) -> None:
    paragraph = doc.add_paragraph()
    style_paragraph(paragraph, after=after)
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        set_run_font(lead, size=11, bold=True, color=NAVY)
        rest = paragraph.add_run(text[len(bold_lead):])
        set_run_font(rest, size=11, color=color)
    else:
        run = paragraph.add_run(text)
        set_run_font(run, size=11, color=color)


def add_bullet(doc, text: str, *, bold_lead: str | None = None,
               keep_with_next: bool = False) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    style_paragraph(paragraph, after=4, line_spacing=1.25)
    paragraph.paragraph_format.keep_with_next = keep_with_next
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        set_run_font(lead, size=11, bold=True, color=NAVY)
        rest = paragraph.add_run(text[len(bold_lead):])
        set_run_font(rest, size=11, color=INK)
    else:
        run = paragraph.add_run(text)
        set_run_font(run, size=11, color=INK)


def add_callout(doc, label: str, text: str, *, fill: str = LIGHT_BLUE,
                accent: str = BLUE, text_color: str = INK) -> None:
    paragraph = doc.add_paragraph()
    style_paragraph(paragraph, before=2, after=10, line_spacing=1.25)
    paragraph.paragraph_format.left_indent = Pt(12)
    paragraph.paragraph_format.right_indent = Pt(12)
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(12)
    shade_paragraph(paragraph, fill)
    set_paragraph_borders(paragraph, color=accent, left=True)
    lead = paragraph.add_run(f"{label}  ")
    set_run_font(lead, size=11, bold=True, color=accent)
    body = paragraph.add_run(text)
    set_run_font(body, size=11, color=text_color)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("商务方案  ·  第 ")
    set_run_font(run, size=9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, value, end])
    tail = paragraph.add_run(" 页")
    set_run_font(tail, size=9, color=MUTED)


def configure_document(doc: Document) -> None:
    doc.settings.odd_and_even_pages_header_footer = False
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, BLUE_DARK, 10, 5),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    bullet = styles["List Bullet"]
    bullet.font.name = "Arial"
    bullet._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    bullet.font.size = Pt(11)
    bullet.paragraph_format.left_indent = Inches(0.375)
    bullet.paragraph_format.first_line_indent = Inches(-0.188)
    bullet.paragraph_format.space_after = Pt(4)
    bullet.paragraph_format.line_spacing = 1.25

    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    style_paragraph(footer_paragraph, after=0, line_spacing=1)
    add_page_number(footer_paragraph)


def build_document() -> Document:
    doc = Document()
    configure_document(doc)

    kicker = doc.add_paragraph()
    style_paragraph(kicker, before=4, after=4, line_spacing=1)
    run = kicker.add_run("微信订单机器人 · 商务方案")
    set_run_font(run, size=10, bold=True, color=BLUE)

    title = doc.add_paragraph()
    style_paragraph(title, after=6, line_spacing=1.05)
    run = title.add_run("部署方案及收费标准")
    set_run_font(run, size=27, bold=True, color=NAVY)

    subtitle = doc.add_paragraph()
    style_paragraph(subtitle, after=4, line_spacing=1.15)
    run = subtitle.add_run("适用于销售订单群聊录单功能的选型、报价与部署沟通")
    set_run_font(run, size=12.5, color=MUTED)

    meta = doc.add_paragraph()
    style_paragraph(meta, after=18, line_spacing=1)
    run = meta.add_run("商务版  |  2026 年 7 月")
    set_run_font(run, size=9.5, color=MUTED)

    add_callout(
        doc,
        "一分钟看懂",
        "群聊录单属于收费功能。费用由“订单识别额度”与“机器人部署费”两部分组成；"
        "实际选择取决于接单群数量、是否需要品牌定制，以及对稳定性和多人协同的要求。",
    )

    add_heading(doc, "01 费用由两部分组成", 1)
    add_body(
        doc,
        "第一部分是订单识别额度，按成功识别并提交到系统的订单条目消耗；"
        "第二部分是机器人部署费，仅在选择专属机器人时按台、按年收取。",
    )

    add_heading(doc, "订单识别额度套餐", 2)
    add_table(
        doc,
        ["套餐", "价格", "基础条目", "赠送条目", "总条目", "折合单价"],
        [
            ["试用包", "¥5,000", "50,000", "—", "50,000", "¥0.10"],
            ["基础包", "¥10,000", "100,000", "42,857", "142,857", "¥0.07"],
            ["A 套餐", "¥30,000", "300,000", "161,538", "461,538", "¥0.065"],
            ["B 套餐", "¥50,000", "500,000", "333,333", "833,333", "¥0.06"],
            ["C 套餐", "¥100,000", "1,000,000", "1,000,000", "2,000,000", "¥0.05"],
        ],
        [1200, 1500, 1650, 1650, 1650, 1710],
        alignments=[
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.RIGHT,
            WD_ALIGN_PARAGRAPH.RIGHT,
            WD_ALIGN_PARAGRAPH.RIGHT,
            WD_ALIGN_PARAGRAPH.RIGHT,
            WD_ALIGN_PARAGRAPH.RIGHT,
        ],
    )
    note = doc.add_paragraph()
    style_paragraph(note, before=0, after=10, line_spacing=1.15)
    run = note.add_run("注：试用包支持试用一周。")
    set_run_font(run, size=9.5, color=MUTED)

    add_heading(doc, "计费规则", 2)
    add_bullet(
        doc,
        "仅对成功识别并提交至系统的订单条目计费；识别失败或未成功提交的不计费。",
        keep_with_next=True,
    )
    add_bullet(
        doc,
        "后台人工录入按 0.5 个条目计费。",
        bold_lead="后台人工录入",
        keep_with_next=True,
    )
    add_bullet(doc, "群聊机器人自动接单按 1 个条目计费。", bold_lead="群聊机器人自动接单")

    add_heading(doc, "02 机器人费用", 1)
    add_body(
        doc,
        "机器人费用与识别额度分开计算。免费公共机器人适合快速验证；"
        "专属机器人更适合正式、稳定、可持续的业务运行。",
    )
    add_table(
        doc,
        ["机器人类型", "收费标准"],
        [
            ["专属机器人（个人微信 / 企业微信）", "¥3,000 / 台 / 年"],
            ["公共微信机器人", "免费（不推荐长期使用）"],
            ["公共企业微信机器人", "免费"],
        ],
        [4200, 5160],
        alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
    )
    add_callout(
        doc,
        "容量提示",
        "单个机器人建议管理 100 个群以内。群数量增加时，可通过增加机器人数量扩容。",
        fill=AMBER,
        accent=AMBER_TEXT,
        text_color=AMBER_TEXT,
    )

    add_heading(doc, "03 按接单群数量选择部署方案", 1)
    add_table(
        doc,
        ["接单群数量", "推荐方案", "选择说明"],
        [
            ["1～100 个群", "专属微信机器人（推荐）\n或公共企业微信机器人", "兼顾品牌与稳定性；也可先低成本快速上线"],
            ["100～200 个群", "专属企业微信机器人（推荐）", "适合企业级稳定运行、多人协同与统一权限管理"],
            ["200 个群以上", "增加机器人数量，采用多机器人部署", "按群规模扩容，降低单机器人负载"],
        ],
        [1800, 3500, 4060],
        alignments=[
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.LEFT,
        ],
    )

    add_heading(doc, "1～100 个群：两种常用选择", 2)
    add_bullet(
        doc,
        "专属微信机器人：使用企业自己的微信账号，可自定义头像与昵称、支持个性化回复，品牌形象统一，独立运行更稳定；费用为 ¥3,000 / 台 / 年。",
        bold_lead="专属微信机器人：",
        keep_with_next=True,
    )
    add_bullet(
        doc,
        "公共企业微信机器人：无需购买机器人，开通即可使用，由企业微信统一管理，适合中小规模企业快速上线。",
        bold_lead="公共企业微信机器人：",
    )

    add_heading(doc, "100～200 个群：优先选择专属企业微信机器人", 2)
    add_body(
        doc,
        "该方案支持企业级稳定运行、多人协同和权限统一管理。现有微信群可直接迁移为企业微信群，"
        "无需重新建群或再次邀请客户，后续增加机器人也更方便。费用为 ¥3,000 / 台 / 年。",
    )

    add_heading(doc, "04 选型边界与注意事项", 1)
    add_heading(doc, "什么时候不建议使用公共微信机器人", 2)
    add_body(
        doc,
        "公共微信机器人虽然可以快速开通，但因多个企业共同使用，不支持企业品牌定制与个性化回复，"
        "稳定性相对较弱，因此只建议用于体验或短期过渡，不建议作为长期正式业务方案。",
    )

    add_heading(doc, "什么时候应直接选择专属方案", 2)
    add_bullet(doc, "需要统一企业品牌形象，包括自定义头像、昵称和回复话术。")
    add_bullet(doc, "群聊录单已进入正式生产使用，对稳定性要求较高。")
    add_bullet(doc, "接单群接近或超过 100 个，需要多人协同、权限管理或后续扩容。")
    add_bullet(doc, "希望将现有微信群迁移为企业微信群，并降低客户重新入群的沟通成本。")

    add_heading(doc, "05 联系观麦人员前，请先确认", 1)
    add_bullet(
        doc,
        "当前接单群数量，以及未来 6～12 个月的预计增长规模。",
        bold_lead="接单群规模：",
        keep_with_next=True,
    )
    add_bullet(
        doc,
        "计划使用公共机器人还是专属机器人；专属方案选择个人微信或企业微信。",
        bold_lead="机器人类型：",
        keep_with_next=True,
    )
    add_bullet(
        doc,
        "希望购买的识别额度套餐，以及预估每月成功提交的订单条目数。",
        bold_lead="额度套餐：",
        keep_with_next=True,
    )
    add_bullet(
        doc,
        "是否需要自定义头像、昵称、个性化回复和品牌统一展示。",
        bold_lead="品牌要求：",
        keep_with_next=True,
    )
    add_bullet(doc, "是否涉及现有微信群迁移、多人协同、权限统一或多机器人扩容。", bold_lead="部署要求：")

    add_callout(
        doc,
        "开通说明",
        "群聊功能为收费功能，请联系观麦人员确认额度套餐、机器人类型与部署数量后开通。",
    )

    add_heading(doc, "06 快速选择", 1)
    add_table(
        doc,
        ["你的情况", "建议选择"],
        [
            ["先验证群聊录单效果", "试用包 + 公共机器人"],
            ["1～100 个群，重视品牌与稳定性", "专属微信机器人"],
            ["1～100 个群，希望快速低成本上线", "公共企业微信机器人"],
            ["100～200 个群，需协同与权限管理", "专属企业微信机器人"],
            ["超过 200 个群", "多机器人部署"],
        ],
        [4680, 4680],
        alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
    )

    return doc


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
